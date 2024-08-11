from docx import  Document
import tkinter
from tkinter import filedialog
from bs4 import BeautifulSoup

#goal of this file is to make a word document from given tags

def get_target_dir() -> str:
    """
    Opens a dialog to select a target directory.

    Returns the path of the selected directory.

    Returns:
    str: The path of the selected directory.11
    """
    tkinter.Tk().withdraw()  # prevents an empty tkinter window from appearing
    folder_path = filedialog.askdirectory()
    return folder_path

def get_tags(text:str)-> list[dict[str, any]]:
    """
    Extracts tags from the given text content and returns a list of dictionaries representing each tag.

    Parameters:
    text (str): The text content containing HTML-like tags.

    Returns:
    list: A list of dictionaries representing tags with keys 'start', 'end', 'structure', 'type', 'attributes', and 'level'.
    """

    tags:list[dict] = []
    #sample tags strcture:
    # tags = [{'start':0, 'end':3, 'structure':'opening', 'type':'p', 'attributes':(r'type="blablabla"'), 'level':0}]

    for index, char in enumerate(text):
        # Extracts information about each tag such as structure, type, attributes, etc.


        #gets opening and closing index (index of < & >)
        if char != '<':
            continue

        closing_index = text.find('>', index + 1)
        
    #gets structure variable
        structure:str | None = None
        if text[index + 1] == r'/': #</tag>
            structure = 'closing'
        elif text[closing_index - 1] == r'/': #<tag/>
            structure = 'self-closing'
        else: #<tag>
            structure = 'opening'

        #get the content
        #first get range of the content in between <>
        #Then using that range gets a string of that content

        content_range:tuple[int,int] = ()  #refers to the content in between the <>
        if structure == 'opening':
            content_range = (index + 1, closing_index)  
        else:
            content_range = (index + 2, closing_index)
            if structure == 'self-closing':
                content_range = (index + 1, closing_index - 1)

        content:str | None = None
        start , end = content_range
        _content:list[str] = []

        for i, v in enumerate(range(start,end)):
            _content.append(text[v])

        content = ''.join(_content)

# get the tag type variable <tag>
        #                    ^^^
        content_separated:list[str] = content.split()
        type = content_separated[0]
        attributes:list[str] = []

#puts attributes into a list of strs
        if len(content_separated) > 1:
            for i,attribute in enumerate(content_separated[1:]):
                attributes.append(attribute)

        tag_dictionary:dict[str,any] = {'start':index, 'end':closing_index, 'structure':structure, 'type':type, 'attributes':attributes}
        tags.append(tag_dictionary)

    level: int = 0
    prev_structure: str|None = None

    for tag in tags:
        """
        ALGORITHM:

        1. Starts at level 0.
        2. If tag structure is same as old structure, level increases by 1.
        3. Else if it is different, level decreases by 1.
        4. If tag structure is self-closing, it treats its structure as if it is the same as the old structure, But returns self-closing to old structure variable.
        5. Finally, if a self-closing structured tag, precedes a closing tag, level decreases by 1. 
        """

        tag_structure = tag.get('structure')
        old_level = level

        def update_level(level:int) -> int:
            if prev_structure == 'closing':
                level -= 1
            if prev_structure == 'opening':
                level += 1
            
            return level

        if tag_structure == 'self-closing':
            # level = update_level(level)
            tag['level'] = level
            # prev_structure = tag_structure

            continue
    
        level = update_level(level)

        if prev_structure == 'self-closing':
            if tag_structure == 'closing':
                level -= 1
            if tag_structure == 'opening':
                level += 1

            tag['level'] = level
            prev_structure = tag_structure

            continue

        if tag_structure != prev_structure :
            level = old_level

        tag['level'] = level
        prev_structure = tag_structure

        continue

    return tags

def pair_tags(tags:list[dict[str, any]]) ->list[tuple[dict[str, any], dict[str, any]|None]]:
    """
    Pair opening and closing tags from a list of tags.

    This function pairs opening tags by finding the next closing tag of same type and level.

    Parameters:
    tags (list): A list of dictionaries representing tags with keys 'structure', 'type', 'level', and 'start'.

    Returns:
    list: A list of tuples containing paired tags where the second element is None for self-closing tags.
    """
    paired_tags:list[tuple[dict,dict]] = []

    _tags:list[dict[str, any]] = tags[:]  # Create a copy of the tags list

    for tag in _tags:
        tag_structure = tag.get('structure')
        tag_type = tag.get('type')
        tag_level = tag.get('level')
        tag_start = tag.get('start')

        if tag_structure == 'self-closing':
            paired_tags.append((tag, None))
            continue

        if tag_structure == 'closing':
            continue

        # Find the next tag with the same type and level
        condition = lambda tag: (
            tag.get('type') == tag_type and
            tag.get('level') == tag_level and
            tag.get('structure') == 'closing' and
            tag.get('start') > tag_start
        )

        closing_tag = next((_tag for _tag in _tags if condition(_tag)), None)

        paired_tags.append((tag, closing_tag))

    pass
    return paired_tags



def get_elements(tags:list[tuple[dict,dict]], text:str) -> list[dict[str, tuple[int, int] | list[id]]]:
    """
    Extracts elements from the given text and their corresponding parent tags.

    Parameters:
    tags (list): A list of dictionaries representing paired tags with keys 'structure', 'type', 'level', and 'start'.
    text (str): The text content from which to extract elements.

    Returns:
    list: A list of dictionaries representing elements with keys 'element', 'element_range', and 'parent_tags'.

    This function processes the text content to extract individual elements and determine their parent tags based on the provided tag pairs. It splits the text into elements using BeautifulSoup, identifies the index ranges for each tag pair, and then extracts the elements and their index ranges. Finally, it assigns the corresponding parent tags to each element based on the tag pair ranges.

    The output is a list of dictionaries, each representing a element with the following keys:
    - 'element': The extracted element content.
    - 'element_range': A tuple containing the start and end indices of the element within the text.
    - 'parent_tags': A list of IDs representing the parent tags associated with the element.
    """

    # Create a new list to hold the combined elements and self-closing tags
    arranged_elements:list[str] = []

    # This block fills in the 'arranged_elements' list
    #made just to avoid repitition
    def node_append(text_body:str)->None:
        bs4_elements_to_search:BeautifulSoup = BeautifulSoup(text_body, 'html.parser')
        text_nodes:list = bs4_elements_to_search.find_all(text=True)
        text_nodes = [text_node.replace('\n','').strip() for text_node in text_nodes if text_node.replace('\n','').strip()]
        arranged_elements.extend(text_nodes)

    search_start:int = 0
    for tag in tags: #finds the nearest self closing tag and uses its location in the xml to  get all the text before it. 
        tag1 = tag[0]
        if tag1.get('structure') !='self-closing':
            continue

        self_closing_tags_string:str = f'<{tag1.get('type')} {' '.join(tag1.get('attributes'))}/>' 

        search_end = tag1.get('start')
        text_to_search = text[search_start:search_end]
        search_start = search_end

        node_append(text_to_search)
        arranged_elements.append(self_closing_tags_string)

    last_string_search:str = text[search_start:]
    node_append(last_string_search)

    # List to store the index ranges for each tag pair
    tag_pair_ranges: list[dict[str, id | tuple[int, int]]] = []

    # Iterate through tag pairs to fill in the tag_pair_ranges dictionary
    for tag_pair in tags:
        range_end: int = None
        range_start: int = None
        _tag_pair_container: dict[str, id | tuple[int, int]] = {}



        if tag_pair[0].get('structure') == 'self-closing':
            continue

        # Determine the start and end index for the tag pair
        if tag_pair[0] and isinstance(tag_pair[0], dict):
            range_start = tag_pair[0].get('end')
        if tag_pair[1] and isinstance(tag_pair[1], dict):
            range_end = tag_pair[1].get('start')

        # Store the index range in a tuple
        _range: tuple[int, int] = (range_start, range_end)
        _tag_pair_container['properties'] = tag_pair
        _tag_pair_container['tag_pair_range'] = _range
        tag_pair_ranges.append(_tag_pair_container)

    # Extracts the elements and their index ranges
    elements_ranges: list[dict[str, tuple[int,int]]] = []
    elements_to_skip:list[str] = []
    _text:str = str(text) #TODO: check if this could cause problems regarding the indexes because it might be different because /n is stripped and removed

    for element in sorted(arranged_elements,key=len, reverse=True):
        if element in elements_to_skip:
            continue
        element_range:dict[str,str|tuple[int,int]] = {}
        occs:int = _text.count(element)
        prev_index:int = 0
        current_index:int = 0
        for _ in range(occs):
            current_index = _text.find(element, current_index + prev_index)
            end_index:int = current_index + len(element)
            _text = list(_text)

            for i in range(current_index, end_index): # do not turn to one liner
                _text[i] = '*'

            _text = ''.join(_text)
            element_range = {'element': element, 'range': (current_index, end_index)}
            elements_ranges.append(element_range)
            prev_index = current_index

        elements_to_skip.append(element)

    elements_ranges = sorted(elements_ranges, key=lambda element_range: element_range.get('range')[0])

    # Add the parent tags for each element range
    for element_range in elements_ranges:
         
        parent_tags: list[id] = []
        element_range_start, element_range_end = element_range.get('range')
        element = element_range.get('element')
        for pair_range in tag_pair_ranges:
             
            tag_range_start, tag_range_end = pair_range.get('tag_pair_range')
            tag_pair_properties = pair_range.get('properties')
             
            if element_range_start > tag_range_start and element_range_end < tag_range_end:
                parent_tags.append(tag_pair_properties)
        element_range['parent_tags'] = parent_tags 

    return elements_ranges 
    # AFTER THIS FUNCTION IS EXECUTED ALL THAT NEEDS TO BE DONE IS FOR EACH ELEMENT TO BE INSTANTIATED AS A

def docElementinstantiator(elements,tags,document):
    #PHASE 1:
    #first separate each element into groups based on its first paragraph parent
    # make sure that elements maintain chronological order
    #NOTE never have nested paragraph

    paragraphTags=  [tag for tag in tags if tag[0].get('type') == 'p'] #puts all paragraph tags in a list

    paragraph_groups = []
    for paragraphTag in paragraphTags:
        pTag_range = (paragraphTag[0].get('start'), paragraphTag[1].get('end'))
        same_parent_paragraph = []
        for element in elements:
            element_range = element.get('range')
            if pTag_range[0] < element_range[0] < pTag_range[1] and pTag_range[1] > element_range[1] > pTag_range[0]: #Basically, if range (w,x) is within (y,z)
                same_parent_paragraph.append(element)
        paragraph_groups.append(same_parent_paragraph)
                

    ...
    #PHASE 2:
    #Once a list has been created for this, instatiate a paragraph for each list then keep that in a dict like so:
    #{paragraph: paragraph_elemnt, elements: [...]}
    #store these dicts in a list
    instantiatedParagraphs = []
    for paragraphGroup in paragraph_groups:
        docParagraph = document.add_paragraph() #FOR FUTURE PURPOSES: if need be add paragraph styles here
        instantiatedParagraphs.append({'paragraph': docParagraph, 'elements': paragraphGroup})


    #PHASE 3:
    #In this phase the actual elements are being created in the file itself
    #assuming that the code for the actual creation is already made (done in a class)
    #The only purpose of this phase is to ensure that they are made in the correct order


    pass

def create_document():
    doc = Document()

    #requests to add in the document
    paragraph = doc.add_paragraph()
    italic = paragraph.add_run('italic')
    italic.italic = True
    bold = paragraph.add_run('bold')
    bold.bold = True

    save_directory = get_target_dir()
    file_name = input('FILE NAME: ') + '.docx'

    doc.save(fr'{save_directory}/{file_name}')

text = '''
<tag1>
    more
    <tag2>
        foo
        <p>
        more
            <testtag1>
             test1
                <testtag1>
                    test2
                </testtag1>
                <testtag1>
                    test3
                    <testtag1>
                        test4
                    </testtag1>
                    test5
                </testtag1>
            </testtag1>
        </p>
    
        <p>
        trick
        </p>
        foo
    </tag2>
    split
</tag1>
'''

TEMPORARYDOCUMENT = Document()

tags = get_tags(text)

tags = pair_tags(tags)

elements = get_elements(tags,text)

docElementinstantiator(elements, tags, TEMPORARYDOCUMENT)

print('done')
